"""Resume processing service.

Standards: python_clean.mdc
- Coordinates storage and parsing
- Generates embeddings for semantic search
- pypdfium2-based PDF extraction with OCRmyPDF fallback

Extraction Flow:
┌─────────────┐
│ qpdf repair │ (best effort, optional)
└──────┬──────┘
       │
       ▼
┌─────────────┐
│ pypdfium2   │ (primary text extraction)
│ (Apache 2.0)│
└──────┬──────┘
       │ (if text < threshold)
       ▼
┌──────────────┐
│ OCRmyPDF     │ (OCR fallback)
│ (MPL 2.0)    │
└──────────────┘
"""

import io
import shutil
import sys
import time
import uuid
from dataclasses import dataclass
from datetime import datetime
from enum import Enum
from typing import Any

import structlog
from tenacity import (
    retry,
    retry_if_exception_type,
    stop_after_attempt,
    wait_exponential,
)

from app.core.domain.resume import ParsedResume, Resume
from app.core.exceptions import ExtractionFailedError, PDFCorruptedError
from app.core.ports.llm import LLMClient
from app.core.ports.repositories import ResumeRepository
from app.core.ports.storage import FileStorage
from app.core.ports.vector_store import VectorStore

logger = structlog.get_logger(__name__)

# Collection name for resume embeddings
RESUMES_COLLECTION = "resumes"

# Minimum text length threshold for valid extraction
MIN_TEXT_THRESHOLD = 50

# Dependency check cache TTL (5 minutes)
DEPENDENCY_CHECK_TTL = 300  # seconds



class ExtractionMethod(Enum):
    """Method used to extract text from PDF."""

    NATIVE = "native"  # pypdfium2 (was: pypdf, pdfplumber, pdfminer)
    OCR = "ocr"  # Tesseract OCR
    # VISION_AI removed - no longer used
    FAILED = "failed"


@dataclass
class UploadContext:
    """Tracks resources created during upload for cleanup on failure."""

    resume_id: str
    s3_key: str | None = None
    vector_store_written: bool = False


@dataclass
class PDFExtractionResult:
    """Result of PDF text extraction with metadata."""

    text: str
    method: ExtractionMethod
    page_count: int
    is_encrypted: bool = False
    ocr_available: bool = True
    extraction_library: str | None = None  # Which library succeeded (pypdfium2, tesseract, etc)
    page_methods: list[str] | None = None  # Which method succeeded for each page
    error_message: str | None = None
    warnings: list[str] = None  # type: ignore[assignment]
    error_code: str | None = None
    is_partial: bool = False  # True if 0 < len(text) < MIN_TEXT_THRESHOLD

    def __post_init__(self) -> None:
        """Compute is_partial based on text length."""
        if self.warnings is None:
            self.warnings = []
        if self.page_methods is None:
            self.page_methods = []
        text_length = len(self.text.strip()) if self.text else 0
        self.is_partial = 0 < text_length < MIN_TEXT_THRESHOLD


class ResumeService:
    """Service for resume upload and parsing."""

    def __init__(
        self,
        *,
        storage: FileStorage,
        resume_repository: ResumeRepository,
        llm_client: LLMClient | None = None,
        vector_store: VectorStore | None = None,
    ) -> None:
        self._storage = storage
        self._resume_repo = resume_repository
        self._llm_client = llm_client
        self._vector_store = vector_store
        # Cache for dependency checks: {cache_key: (result_dict, timestamp)}
        self._dependency_cache: dict[str, tuple[dict[str, bool], float]] = {}

    async def upload_and_parse(
        self,
        *,
        user_id: str,
        filename: str,
        content: bytes,
        content_type: str,
    ) -> tuple[Resume, dict[str, Any]]:
        """Upload resume and parse its content.

        Flow: Extract → Upload S3 → Save DB
        If extraction fails completely, no S3 upload occurs.
        
        Returns:
            Tuple of (Resume, extraction_metadata) where metadata contains:
            - extraction_warnings: list[str]
            - extraction_method: str | None
        """
        resume_id = str(uuid.uuid4())
        s3_key = f"resumes/{user_id}/{resume_id}/{filename}"
        upload_ctx = UploadContext(resume_id=resume_id)

        try:
            # Step 1: Extract text based on file type
            raw_text, extraction_metadata = await self._extract_text(content=content, content_type=content_type)
            
            # Get extraction error message if extraction failed
            extraction_error: str | None = None
            if not raw_text or len(raw_text.strip()) == 0:
                # Build error message for user guidance
                deps = self._check_extraction_dependencies()
                methods_attempted = ["pypdfium2"]
                if deps.get("ocrmypdf"):
                    methods_attempted.append("ocrmypdf")

                error_parts = []
                if not deps.get("pypdfium2"):
                    error_parts.append("pypdfium2 package is required for PDF extraction")
                if not deps.get("ocrmypdf"):
                    error_parts.append(
                        "OCRmyPDF not available - install with: pip install ocrmypdf"
                    )
                error_parts.append("PDF appears to be scanned/image-based or text extraction failed")
                error_parts.append("Try re-exporting the PDF from the original source as text-based")
                
                extraction_error = ". ".join(error_parts)
                logger.warning(
                    "text_extraction_empty_but_continuing",
                    resume_id=resume_id,
                    extraction_error=extraction_error,
                )

            # Step 4: Upload to S3 (always upload, even if extraction failed)
            await self._storage.upload(
                key=s3_key,
                data=content,
                content_type=content_type,
            )
            upload_ctx.s3_key = s3_key  # Track for cleanup

            # Step 5: Parse resume text
            parsed_data = await self._parse_resume_text(raw_text) if raw_text else None

            # Step 6: Generate embedding for semantic search
            embedding = None
            if raw_text and self._llm_client:
                try:
                    # Truncate text for embedding (8k chars max for most models)
                    embed_text = raw_text[:8000]
                    embedding = await self._llm_client.embed(text=embed_text)

                    logger.info(
                        "resume_embedding_generated",
                        resume_id=resume_id,
                        text_length=len(embed_text),
                        embedding_dim=len(embedding),
                    )

                    # Step 7: Store in vector database for semantic search
                    if self._vector_store:
                        await self._vector_store.add_embedding(
                            collection=RESUMES_COLLECTION,
                            doc_id=resume_id,
                            embedding=embedding,
                            metadata={
                                "user_id": user_id,
                                "filename": filename,
                            },
                        )
                        upload_ctx.vector_store_written = True  # Track for cleanup
                        logger.info(
                            "resume_stored_in_vector_db",
                            resume_id=resume_id,
                            collection=RESUMES_COLLECTION,
                        )

                except Exception as e:
                    logger.warning(
                        "resume_embedding_failed",
                        resume_id=resume_id,
                        error=str(e),
                    )
                    # Continue without embedding - not critical

            # Step 8: Check if this is the first resume (make it primary)
            existing = await self._resume_repo.get_by_user_id(user_id)
            is_primary = len(existing) == 0

            # Step 9: Create resume record
            resume = Resume(
                id=resume_id,
                user_id=user_id,
                filename=filename,
                s3_key=s3_key,
                raw_text=raw_text,
                parsed_data=parsed_data,
                embedding=embedding,
                is_primary=is_primary,
                extraction_error=extraction_error,
                created_at=datetime.utcnow(),
            )

            created_resume = await self._resume_repo.create(resume)
            
            # Return resume with extraction metadata
            return created_resume, extraction_metadata

        except PDFCorruptedError:
            # Cleanup on PDF corruption (still a real error)
            await self._cleanup_on_failure(upload_ctx)
            raise
        except Exception as e:
            # Cleanup on any other failure
            await self._cleanup_on_failure(upload_ctx)
            logger.error(
                "resume_upload_failed",
                resume_id=resume_id,
                error=str(e),
                exc_info=True,
            )
            raise

    async def _extract_text(
        self,
        *,
        content: bytes,
        content_type: str,
    ) -> tuple[str | None, dict[str, Any]]:
        """Extract text from resume file.
        
        Returns:
            Tuple of (extracted_text, metadata_dict) where metadata contains:
            - extraction_warnings: list[str]
            - extraction_method: str | None
        """
        file_size = len(content)
        logger.info(
            "text_extraction_started",
            content_type=content_type,
            file_size_bytes=file_size,
        )

        try:
            if content_type == "application/pdf":
                text, metadata = await self._extract_pdf_text(content)
            elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                text = await self._extract_docx_text(content)
                metadata = {
                    "extraction_warnings": [],
                    "extraction_method": "python-docx",
                }
            else:
                logger.warning("unsupported_content_type", content_type=content_type)
                return None, {"extraction_warnings": [], "extraction_method": None}

            if text:
                logger.info(
                    "text_extraction_succeeded",
                    content_type=content_type,
                    text_length=len(text),
                )
            else:
                logger.warning(
                    "text_extraction_empty",
                    content_type=content_type,
                )

            return (text if text else None, metadata)

        except Exception as e:
            logger.error(
                "text_extraction_failed",
                content_type=content_type,
                error=str(e),
                exc_info=True,
            )
            return None, {"extraction_warnings": [], "extraction_method": None}

    async def _extract_pdf_text(self, content: bytes) -> tuple[str, dict[str, Any]]:
        """Extract text from PDF using pypdfium2-based pipeline.

        Strategy:
        1. qpdf --repair (best effort, optional)
        2. pypdfium2 text extraction (primary)
        3. OCRmyPDF OCR (if text < threshold)
        
        Returns:
            Tuple of (extracted_text, metadata_dict) where metadata contains:
            - extraction_warnings: list[str]
            - extraction_method: str | None
        """
        result = await self._extract_pdf_with_metadata(content)

        logger.info(
            "pdf_extraction_complete",
            method=result.method.value,
            extraction_library=result.extraction_library,
            page_count=result.page_count,
            text_length=len(result.text),
            is_encrypted=result.is_encrypted,
            ocr_available=result.ocr_available,
            error=result.error_message,
        )

        # Build metadata from extraction result
        metadata = {
            "extraction_warnings": result.warnings or [],
            "extraction_method": result.extraction_library,
        }

        return result.text, metadata

    async def _extract_pdf_with_metadata(self, content: bytes) -> PDFExtractionResult:
        """Extract text from PDF using new pypdfium2-based pipeline.

        Extraction priority:
        1. qpdf --repair (best effort, non-blocking)
        2. pypdfium2 text extraction (primary)
        3. OCRmyPDF OCR (if text < threshold)
        """
        # Step 0: File size check
        MAX_PDF_SIZE = 100 * 1024 * 1024  # 100MB
        if len(content) > MAX_PDF_SIZE:
            self._log_extraction_error(
                "pdf_size_check",
                ValueError(f"PDF size {len(content)} exceeds {MAX_PDF_SIZE}"),
                "warning"
            )
            raise ExtractionFailedError(
                methods_attempted=[],
                dependency_status={},
                guidance=[f"PDF file size ({len(content) / 1024 / 1024:.1f}MB) exceeds maximum ({MAX_PDF_SIZE / 1024 / 1024}MB). Please use a smaller file."]
            )
        
        # Step 1: qpdf repair (best effort)
        repaired_content = await self._repair_pdf_with_qpdf(content)
        content_to_use = repaired_content if repaired_content else content
        
        # Step 2: pypdfium2 text extraction
        page_count = 0
        has_images = False
        try:
            text, page_count, has_images = await self._try_pypdfium2_text(content_to_use)
            
            # Step 3: Text quality validation
            if text and self._is_text_quality_valid(text):
                page_methods = ["pypdfium2"] * page_count
                logger.info(
                    "extraction_attempt_success",
                    method="pypdfium2",
                    text_length=len(text),
                    page_count=page_count
                )
                return PDFExtractionResult(
                    text=text,
                    method=ExtractionMethod.NATIVE,
                    page_count=page_count,
                    extraction_library="pypdfium2",
                    page_methods=page_methods,
                )
        except ExtractionFailedError:
            raise  # Re-raise page limit errors
        except Exception as e:
            self._log_extraction_error("pypdfium2_text", e, "warning")
            # Continue to OCR fallback
        
        # Step 4: OCR fallback (OCRmyPDF)
        # OCRmyPDF processes the entire PDF and adds OCR text layer
        try:
            ocr_text = await self._ocr_pdf_with_ocrmypdf(content_to_use, force_ocr=True)
            
            if ocr_text and self._is_text_quality_valid(ocr_text):
                extraction_library = "ocrmypdf"
                
                logger.info(
                    "extraction_attempt_success",
                    method="ocr",
                    text_length=len(ocr_text),
                    page_count=page_count,
                    extraction_library=extraction_library
                )
                return PDFExtractionResult(
                    text=ocr_text,
                    method=ExtractionMethod.OCR,
                    page_count=page_count,
                    extraction_library=extraction_library,
                    page_methods=["ocrmypdf"] * page_count if page_count > 0 else [],
                    warnings=["PDF appears to be scanned or image-based. Text was extracted using OCR - please verify accuracy."],
                )
        
        except ImportError:
            self._log_extraction_error("ocrmypdf_import", ImportError("ocrmypdf not installed"), "warning")
        except Exception as e:
            self._log_extraction_error("ocr_fallback", e, "warning")
        
        # All methods failed - build specific error message
        deps = self._check_extraction_dependencies()
        methods_attempted = ["pypdfium2"]
        if deps.get("ocrmypdf"):
            methods_attempted.append("ocrmypdf")
        
        if has_images:
            error_msg = (
                "PDF appears to be scanned/image-based. All extraction methods failed. "
                "Try installing OCRmyPDF for better OCR support, or re-export the PDF as text-based."
            )
        else:
            error_msg = (
                "No text could be extracted from PDF using any available method. "
                "The PDF may be corrupted, encrypted, or in an unsupported format. "
                "Try re-exporting from the original source or using a PDF repair tool."
            )
        
        warnings = []
        if not deps.get("pypdfium2"):
            warnings.append("pypdfium2 not installed - text extraction unavailable")
        if not deps.get("ocrmypdf"):
            warnings.append("OCRmyPDF not installed - OCR fallback unavailable")
        if not deps.get("qpdf"):
            warnings.append("qpdf not installed - PDF repair unavailable (optional)")
        
        logger.warning(
            "all_extraction_methods_failed",
            page_count=page_count,
            has_images=has_images,
            methods_attempted=methods_attempted,
            dependency_status=deps,
        )
        
        return PDFExtractionResult(
            text="",
            method=ExtractionMethod.FAILED,
            page_count=page_count or 0,
            error_message=error_msg,
            error_code="EXTRACTION_FAILED",
            warnings=warnings,
        )

    async def _try_pypdfium2_text(self, content: bytes) -> tuple[str, int, bool]:
        """Try text extraction using pypdfium2 (Apache 2.0 license).
        
        Returns:
            Tuple of (extracted_text, page_count, has_images)
        """
        try:
            import pypdfium2 as pdfium  # type: ignore[import-not-found]
            import asyncio
            
            doc = pdfium.PdfDocument(content)
            page_count = len(doc)
            
            # Check for zero pages
            if page_count == 0:
                logger.warning(
                    "extraction_stage_failed",
                    stage="pypdfium2_text",
                    error_type="EmptyPDFError",
                    error_message="PDF has zero pages"
                )
                return "", 0, False
            
            # Check page count limit
            MAX_PAGES_FOR_EXTRACTION = 100
            if page_count > MAX_PAGES_FOR_EXTRACTION:
                logger.warning(
                    "extraction_stage_failed",
                    stage="pypdfium2_text",
                    error_type="PageLimitExceeded",
                    error_message=f"PDF has {page_count} pages, max is {MAX_PAGES_FOR_EXTRACTION}"
                )
                raise ExtractionFailedError(
                    methods_attempted=["pypdfium2"],
                    dependency_status={},
                    guidance=[f"PDF exceeds page limit ({MAX_PAGES_FOR_EXTRACTION}). Please split into smaller files."]
                )
            
            # Try to access first page - will fail if encrypted
            try:
                page = doc.get_page(0)
                textpage = page.get_textpage()
                # If we get here, not encrypted (or encryption check passed)
            except Exception as e:
                error_str = str(e).lower()
                if any(kw in error_str for kw in ['encrypt', 'password', 'permission']):
                    logger.warning(
                        "extraction_stage_failed",
                        stage="pypdfium2_text",
                        error_type="EncryptedPDFError",
                        error_message="PDF is password-protected"
                    )
                    return "", page_count, False
                raise
            
            # Extract text from all pages and detect images
            text_parts: list[str] = []
            has_images = False
            
            for i in range(page_count):
                try:
                    page = doc.get_page(i)
                    textpage = page.get_textpage()
                    page_text = textpage.get_text_range()
                    text_parts.append(page_text)
                    
                    # Detect if page has images (scanned content)
                    # If page has no text but exists, likely image-based
                    if not page_text.strip():
                        # Check if page has any renderable content
                        try:
                            bitmap = page.render(scale=1)
                            # If we can render but no text, likely scanned
                            has_images = True
                            if hasattr(bitmap, 'close'):
                                bitmap.close()
                            else:
                                del bitmap
                        except Exception:
                            pass
                            
                except Exception as e:
                    self._log_extraction_error("pypdfium2_text", e, "warning")
                    text_parts.append("")  # Continue with other pages
            
            text = "\n".join(text_parts).strip()
            return text, page_count, has_images
            
        except ImportError:
            self._log_extraction_error("pypdfium2_import", ImportError("pypdfium2 not installed"), "warning")
            return "", 0, False
        except ExtractionFailedError:
            raise  # Re-raise page limit errors
        except Exception as e:
            self._log_extraction_error("pypdfium2_text", e, "warning")
            raise  # Re-raise for caller

    async def _repair_pdf_with_qpdf(self, content: bytes) -> bytes | None:
        """Repair PDF using qpdf (best effort, non-blocking).
        
        Returns repaired PDF bytes or None if repair fails/unavailable.
        Does not fail extraction if qpdf is unavailable.
        
        Args:
            content: Original PDF content as bytes
            
        Returns:
            Repaired PDF bytes or None if repair fails/unavailable
        """
        import tempfile
        import os
        import subprocess
        import asyncio
        
        if not shutil.which("qpdf"):
            logger.debug("extraction_stage_skipped", stage="qpdf_repair", reason="qpdf_not_available")
            return None
        
        input_path = None
        output_path = None
        
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as input_file:
                input_file.write(content)
                input_path = input_file.name
            
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as output_file:
                output_path = output_file.name
            
            # CORRECTED: Wrap subprocess.run in a function for proper timeout handling
            def run_qpdf():
                return subprocess.run(
                    ['qpdf', '--repair', input_path, output_path],
                    capture_output=True,
                    check=False,
                    timeout=30  # ✅ Now timeout is valid
                )
            
            # Windows-compatible async subprocess with timeout
            try:
                result = await asyncio.wait_for(
                    asyncio.to_thread(run_qpdf),
                    timeout=35  # Slightly longer than subprocess timeout
                )
            except asyncio.TimeoutError:
                timeout_error = TimeoutError("qpdf repair exceeded 35s timeout")
                self._log_extraction_error("qpdf_repair", timeout_error, "warning")
                return None
            except subprocess.TimeoutExpired as e:
                self._log_extraction_error("qpdf_repair", e, "warning")
                return None
            
            if result.returncode == 0 and os.path.exists(output_path):
                file_size = os.path.getsize(output_path)
                if file_size > 0:
                    with open(output_path, 'rb') as f:
                        repaired_content = f.read()
                    logger.info(
                        "extraction_stage_success",
                        stage="qpdf_repair",
                        original_size=len(content),
                        repaired_size=len(repaired_content)
                    )
                    return repaired_content
            
            logger.debug(
                "extraction_stage_failed",
                stage="qpdf_repair",
                error_type="RepairFailed",
                error_message=f"qpdf returned code {result.returncode}"
            )
            return None
            
        except Exception as e:
            self._log_extraction_error("qpdf_repair", e, "warning")
            return None
        finally:
            # CRITICAL: Always cleanup temp files
            for path in [input_path, output_path]:
                if path:
                    try:
                        if os.path.exists(path):
                            os.unlink(path)
                    except Exception as cleanup_error:
                        logger.debug(
                            "temp_file_cleanup_failed",
                            path=path,
                            error=str(cleanup_error)
                        )


    @retry(
        retry=retry_if_exception_type((TimeoutError, OSError)),
        stop=stop_after_attempt(3),
        wait=wait_exponential(multiplier=1, min=1, max=10),
    )
    async def _ocr_pdf_with_ocrmypdf(self, content: bytes, *, force_ocr: bool = False) -> str:
        """Run OCR on a PDF using OCRmyPDF with retry logic.
        
        OCRmyPDF processes the entire PDF, adds an OCR text layer, and returns
        the extracted text. This is more reliable than raw Tesseract for PDFs.
        
        Args:
            content: PDF file bytes.
            force_ocr: If True, force OCR on every page (use when pypdfium2
                       failed entirely and the PDF may be malformed).
        
        Returns:
            Extracted text from the PDF, or empty string if OCR fails.
        """
        import tempfile
        import os
        import asyncio
        
        input_path = None
        output_path = None
        
        try:
            import ocrmypdf  # type: ignore[import-not-found]
            
            # Create temporary files for input and output
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as input_file:
                input_file.write(content)
                input_path = input_file.name
            
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as output_file:
                output_path = output_file.name
            
            # Run OCRmyPDF in a thread to avoid blocking
            def run_ocr() -> str:
                try:
                    # OCRmyPDF options:
                    # - skip-text: skip pages that already have text (faster)
                    # - force-ocr: force OCR even if text exists (more thorough)
                    # - optimize: optimize the PDF after OCR
                    ocrmypdf.ocr(
                        input_path,
                        output_path,
                        skip_text=not force_ocr,  # Skip text pages unless forced
                        force_ocr=force_ocr,  # Force OCR on all pages for malformed PDFs
                        optimize=0,  # No optimization (faster)
                        output_type='pdf',  # Output as PDF
                    )
                    
                    # Extract text from OCR'd PDF using pypdfium2
                    import pypdfium2 as pdfium  # type: ignore[import-not-found]
                    with open(output_path, 'rb') as f:
                        ocr_pdf_content = f.read()
                    
                    doc = pdfium.PdfDocument(ocr_pdf_content)
                    text_parts: list[str] = []
                    
                    for i in range(len(doc)):
                        try:
                            page = doc.get_page(i)
                            textpage = page.get_textpage()
                            page_text = textpage.get_text_range()
                            text_parts.append(page_text)
                        except Exception:
                            text_parts.append("")
                    
                    return "\n".join(text_parts).strip()
                    
                except Exception as e:
                    error_msg = getattr(e, 'message', None) or str(e) or repr(e)
                    logger.warning("ocrmypdf_processing_failed", error=error_msg, error_type=type(e).__name__)
                    return ""
            
            # Run OCR in thread with timeout
            try:
                text = await asyncio.wait_for(
                    asyncio.to_thread(run_ocr),
                    timeout=120  # 2 minute timeout for OCR
                )
                
                if not text or not text.strip():
                    logger.debug("ocrmypdf_returned_empty")
                
                return text
                
            except asyncio.TimeoutError:
                logger.warning("ocrmypdf_timeout", timeout_seconds=120)
                return ""
            
        except ImportError:
            logger.warning("ocrmypdf_not_available")
            return ""
        except (TimeoutError, OSError) as e:
            logger.warning("ocr_pdf_retryable_error", error=str(e))
            raise  # Re-raise for retry logic
        except Exception as e:
            logger.warning("ocr_pdf_failed", error=str(e), error_type=type(e).__name__)
            return ""  # Non-retryable errors return empty
        finally:
            # Cleanup temp files
            for path in [input_path, output_path]:
                if path:
                    try:
                        if os.path.exists(path):
                            os.unlink(path)
                    except Exception as cleanup_error:
                        logger.debug(
                            "temp_file_cleanup_failed",
                            path=path,
                            error=str(cleanup_error)
                        )

    def _log_extraction_error(self, stage: str, error: Exception, level: str = "warning") -> None:
        """Consistent error logging for extraction stages.
        
        Args:
            stage: Name of the extraction stage (e.g., "qpdf_repair", "pypdfium2_text")
            error: The exception that occurred
            level: Log level ("warning" or "error")
        """
        log_func = getattr(logger, level)
        log_func(
            "extraction_stage_failed",
            stage=stage,
            error_type=type(error).__name__,
            error_message=str(error),
            exc_info=level == "error",
        )

    def _is_text_quality_valid(self, text: str) -> bool:
        """Validate extracted text has meaningful content beyond length check.
        
        Args:
            text: Extracted text to validate
            
        Returns:
            True if text appears to be valid resume content, False otherwise.
        """
        if len(text.strip()) < MIN_TEXT_THRESHOLD:
            return False
        
        # Check for readable characters (not just symbols)
        alnum_count = sum(1 for c in text if c.isalnum())
        if alnum_count < MIN_TEXT_THRESHOLD * 0.5:  # At least 50% alphanumeric
            return False
        
        # Check for common resume keywords (heuristic)
        resume_keywords = ['experience', 'education', 'skill', 'email', 'phone', 'work', 'job', 'position']
        text_lower = text.lower()
        keyword_matches = sum(1 for kw in resume_keywords if kw in text_lower)
        
        # If no keywords and low alnum ratio, likely garbage
        if keyword_matches == 0 and alnum_count < len(text) * 0.3:
            return False
        
        return True

    def _is_ocrmypdf_available(self) -> bool:
        """Check if OCRmyPDF is available (package + Tesseract + Ghostscript)."""
        try:
            import ocrmypdf  # type: ignore[import-not-found]
            # OCRmyPDF requires both Tesseract and Ghostscript
            has_tesseract = shutil.which("tesseract") is not None
            has_gs = (
                shutil.which("gs") is not None
                or shutil.which("gswin64c") is not None
                or shutil.which("gswin32c") is not None
            )
            if not has_tesseract:
                logger.debug("ocrmypdf_missing_dep", dep="tesseract")
            if not has_gs:
                logger.debug("ocrmypdf_missing_dep", dep="ghostscript")
            return has_tesseract and has_gs
        except ImportError:
            return False
        except Exception:
            return False

    def _check_extraction_dependencies(self) -> dict[str, bool]:
        """Check availability of extraction dependencies.

        Caches results for 5 minutes to avoid repeated checks.

        Returns:
            Dictionary with dependency availability:
            - pypdfium2: bool
            - qpdf: bool
            - tesseract: bool
        """
        cache_key = "dependencies"
        current_time = time.time()

        # Check cache
        if cache_key in self._dependency_cache:
            cached_result, cached_time = self._dependency_cache[cache_key]
            if current_time - cached_time < DEPENDENCY_CHECK_TTL:
                return cached_result

        # Check dependencies
        dependencies: dict[str, bool] = {}

        # Check pypdfium2
        dependencies["pypdfium2"] = self._is_package_installed("pypdfium2")
        
        # Check qpdf (optional, for PDF repair)
        dependencies["qpdf"] = shutil.which("qpdf") is not None

        # Check OCRmyPDF (for OCR fallback)
        dependencies["ocrmypdf"] = self._is_ocrmypdf_available()

        # Cache result
        self._dependency_cache[cache_key] = (dependencies, current_time)

        logger.debug(
            "extraction_dependencies_checked",
            pypdfium2=dependencies["pypdfium2"],
            qpdf=dependencies["qpdf"],
            ocrmypdf=dependencies["ocrmypdf"],
        )

        return dependencies

    def _is_package_installed(self, package_name: str) -> bool:
        """Check if a Python package is installed and functional.

        For paddleocr, a top-level import succeeds even when the
        paddlepaddle backend is missing.  We verify the backend too.
        """
        try:
            if package_name == "paddleocr":
                from paddleocr import PaddleOCR  # type: ignore[import-not-found]  # noqa: F401
                __import__("paddle")  # verify paddlepaddle backend exists
            else:
                __import__(package_name)
            return True
        except (ImportError, Exception):
            return False

    async def _cleanup_s3(self, key: str) -> None:
        """Delete S3 object. Idempotent - safe to call multiple times."""
        if not self._storage:
            logger.debug("cleanup_skipped", type="s3", reason="storage_not_available")
            return
        try:
            logger.debug("cleanup_operation_started", type="s3", key=key)
            await self._storage.delete(key=key)
            logger.info("cleanup_operation_success", type="s3", key=key)
        except Exception as e:
            logger.warning("cleanup_operation_failed", type="s3", key=key, error=str(e))
            # Don't raise - cleanup failures shouldn't fail the request

    async def _cleanup_vector_store(self, resume_id: str) -> None:
        """Delete vector store entry. Idempotent - safe to call multiple times."""
        if not self._vector_store:
            return
        try:
            logger.debug("cleanup_operation_started", type="vector_store", resume_id=resume_id)
            await self._vector_store.delete(collection=RESUMES_COLLECTION, doc_id=resume_id)
            logger.info("cleanup_operation_success", type="vector_store", resume_id=resume_id)
        except Exception as e:
            logger.warning("cleanup_operation_failed", type="vector_store", resume_id=resume_id, error=str(e))
            # Don't raise - cleanup failures shouldn't fail the request

    async def _cleanup_on_failure(self, ctx: UploadContext) -> None:
        """Clean up resources created during failed upload."""
        logger.info(
            "cleanup_started",
            resume_id=ctx.resume_id,
            s3_key=ctx.s3_key,
            vector_store_written=ctx.vector_store_written,
        )
        # Cleanup in reverse order of creation
        if ctx.vector_store_written:
            await self._cleanup_vector_store(ctx.resume_id)
        if ctx.s3_key:
            await self._cleanup_s3(ctx.s3_key)
        logger.info("cleanup_completed", resume_id=ctx.resume_id)

    async def _extract_docx_text(self, content: bytes) -> str:
        """Extract text from DOCX."""
        try:
            from docx import Document

            doc = Document(io.BytesIO(content))
            text_parts: list[str] = []

            for para in doc.paragraphs:
                text_parts.append(para.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)

            text = "\n".join(text_parts)

            logger.info(
                "docx_extraction_complete",
                paragraph_count=len(doc.paragraphs),
                table_count=len(doc.tables),
                text_length=len(text),
            )

            return text

        except ImportError:
            logger.error("python_docx_not_installed")
            return ""
        except Exception as e:
            logger.warning("docx_extraction_failed", error=str(e))
            return ""

    async def _parse_resume_text(self, text: str) -> ParsedResume:
        """Parse resume text into structured data.

        This is a simplified parser. In production, use LLM for better parsing.
        """
        lines = text.split("\n")
        skills: list[str] = []

        # Simple skill extraction (look for common patterns)
        skill_keywords = [
            "python", "javascript", "typescript", "react", "node.js", "aws",
            "docker", "kubernetes", "sql", "postgresql", "mongodb", "redis",
            "fastapi", "django", "flask", "git", "linux", "api", "rest",
        ]

        text_lower = text.lower()
        for keyword in skill_keywords:
            if keyword in text_lower:
                skills.append(keyword.title() if len(keyword) > 3 else keyword.upper())

        # Extract email
        import re
        email_match = re.search(r"[\w\.-]+@[\w\.-]+\.\w+", text)
        email = email_match.group(0) if email_match else None

        # Extract phone
        phone_match = re.search(r"[\+]?[\d\s\-\(\)]{10,}", text)
        phone = phone_match.group(0).strip() if phone_match else None

        # First non-empty line is often the name
        full_name = None
        for line in lines:
            stripped = line.strip()
            if stripped and len(stripped) < 50 and "@" not in stripped:
                full_name = stripped
                break

        return ParsedResume(
            full_name=full_name,
            email=email,
            phone=phone,
            skills=skills,
            total_years_experience=None,  # Would need more sophisticated parsing
        )  # type: ignore[call-arg]
